"""
pip install pandas numpy

# PDF处理

pip install PyPDF2 pdfplumber pypdf

# Word处理

pip install python-docx

# 文本处理与评估

pip install nltk langchain tiktoken

# 可选：OCR支持（扫描版PDF）

pip install pytesseract pdf2image"""

import os

import re

import json

from pathlib import Path

from typing import List, Dict, Tuple, Optional

from dataclasses import dataclass, field

from enum import Enum

# PDF处理

import PyPDF2

import pdfplumber

from pypdf import PdfReader

# Word处理

from docx import Document

# 文本处理

import nltk

from nltk.tokenize import sent_tokenize

import tiktoken

# 确保下载NLTK数据

try:

    nltk.data.find('tokenizers/punkt')

except LookupError:

    nltk.download('punkt')



# 可选：用于高级分块

from langchain.text_splitter import (

    RecursiveCharacterTextSplitter,

    NLTKTextSplitter,

    TokenTextSplitter

)



class DocumentType(Enum):

    """文档类型枚举"""

    PDF = "pdf"

    DOCX = "docx"

    TXT = "txt"



@dataclass

class Document:

    """文档数据结构"""

    content: str

    metadata: Dict[str, any]

    pages: Optional[List[str]] = None



class DocumentLoader:

    """统一的文档加载器"""

    @staticmethod

    def load_pdf(file_path: str, method: str = "pdfplumber") -> Document:

        """

       加载PDF文件

        Args:

            file_path: PDF文件路径

            method: 解析方法 ('pdfplumber', 'pypdf2', 'pypdf')

            Returns:

            Document对象

        """

        content_parts = []

        pages_content = []

        metadata = {}

        if method == "pdfplumber":

            with pdfplumber.open(file_path) as pdf:

                metadata = {

                    "pages": len(pdf.pages),

                    "file_name": Path(file_path).name,

                    "method": method

                }

                

                for i, page in enumerate(pdf.pages):

                    text = page.extract_text()

                    if text:

                        pages_content.append(text)

                        content_parts.append(text)

                    else:

                        print(f"警告: 第{i+1}页无法提取文本")

                        

        elif method == "pypdf2":

            with open(file_path, 'rb') as file:

                pdf_reader = PyPDF2.PdfReader(file)

                metadata = {

                    "pages": len(pdf_reader.pages),

                    "file_name": Path(file_path).name,

                    "method": method

                }

                

                for i, page in enumerate(pdf_reader.pages):

                    text = page.extract_text()

                    if text:

                        pages_content.append(text)

                        content_parts.append(text)

                        

        elif method == "pypdf":

            reader = PdfReader(file_path)

            metadata = {

                "pages": len(reader.pages),

                "file_name": Path(file_path).name,

                "method": method

            }

            

            for i, page in enumerate(reader.pages):

                text = page.extract_text()

                if text:

                    pages_content.append(text)

                    content_parts.append(text)

        

        full_content = "\n\n".join(content_parts)

        

        return Document(

            content=full_content,

            metadata=metadata,

            pages=pages_content

        )

    

    @staticmethod

    def load_docx(file_path: str) -> Document:

        """

        加载Word文档

        

        Args:

            file_path: DOCX文件路径

        

        Returns:

            Document对象

        """

        doc = Document(file_path)

        content_parts = []

        

        # 提取段落

        paragraphs = []

        for para in doc.paragraphs:

            if para.text.strip():

                paragraphs.append(para.text)

                content_parts.append(para.text)

        

        # 提取表格

        tables_content = []

        for table in doc.tables:

            table_text = []

            for row in table.rows:

                row_text = []

                for cell in row.cells:

                    if cell.text.strip():

                        row_text.append(cell.text.strip())

                if row_text:

                    table_text.append(" | ".join(row_text))

            if table_text:

                tables_content.append("\n".join(table_text))

                content_parts.extend(tables_content)

        

        metadata = {

            "paragraphs": len(paragraphs),

            "tables": len(doc.tables),

            "file_name": Path(file_path).name,

            "method": "python-docx"

        }

        

        full_content = "\n\n".join(content_parts)

        

        return Document(

            content=full_content,

            metadata=metadata,

            pages=None

        )

    

    @staticmethod

    def load_txt(file_path: str, encoding: str = 'utf-8') -> Document:

        """

        加载纯文本文件

        

        Args:

            file_path: TXT文件路径

            encoding: 文件编码

        

        Returns:

            Document对象

        """

        with open(file_path, 'r', encoding=encoding) as f:

            content = f.read()

        

        metadata = {

            "file_name": Path(file_path).name,

            "size_chars": len(content),

            "method": "txt"

        }

        

        return Document(

            content=content,

            metadata=metadata,

            pages=None

        )

    

    @staticmethod

    def auto_load(file_path: str) -> Document:

        """根据文件扩展名自动选择加载方法"""

        file_ext = Path(file_path).suffix.lower()

        

        if file_ext == '.pdf':

            return DocumentLoader.load_pdf(file_path)

        elif file_ext == '.docx':

            return DocumentLoader.load_docx(file_path)

        elif file_ext == '.txt':

            return DocumentLoader.load_txt(file_path)

        else:

            raise ValueError(f"不支持的文件类型: {file_ext}")



class ChunkingStrategy(Enum):

    """切片策略枚举"""

    FIXED_SIZE = "fixed_size"          # 固定大小

    SENTENCE = "sentence"               # 基于句子

    PARAGRAPH = "paragraph"             # 基于段落

    RECURSIVE = "recursive"             # 递归切片

    SEMANTIC = "semantic"               # 语义切片（基础版）

    OVERLAPPING = "overlapping"         # 重叠切片



@dataclass

class TextChunk:

    """文本块数据结构"""

    id: int

    text: str

    metadata: Dict[str, any]

    start_char: int

    end_char: int

    tokens: Optional[int] = None



class TextChunker:

    """文本切片器"""

    

    def __init__(self, strategy: ChunkingStrategy = ChunkingStrategy.RECURSIVE):

        self.strategy = strategy

        self.chunks: List[TextChunk] = []

        

    def chunk_fixed_size(self, text: str, chunk_size: int = 500, 

                         overlap: int = 0) -> List[TextChunk]:

        """

        固定大小切片

        

        Args:

            text: 输入文本

            chunk_size: 块大小（字符数）

            overlap: 重叠字符数

        """

        chunks = []

        start = 0

        text_length = len(text)

        

        chunk_id = 0

        while start < text_length:

            end = min(start + chunk_size, text_length)

            

            # 尝试在句子边界处截断

            if end < text_length:

                # 查找最近的句子结束符

                for sep in ['. ', '! ', '? ', '\n\n', '\n']:

                    last_sep = text.rfind(sep, start, end)

                    if last_sep != -1:

                        end = last_sep + len(sep)

                        break

            

            chunk_text = text[start:end].strip()

            if chunk_text:

                chunks.append(TextChunk(

                    id=chunk_id,

                    text=chunk_text,

                    metadata={"strategy": "fixed_size", "chunk_size": chunk_size},

                    start_char=start,

                    end_char=end

                ))

                chunk_id += 1

            

            start = end - overlap if overlap > 0 else end

        

        return chunks

    

    def chunk_by_sentences(self, text: str, max_sentences: int = 5,

                          overlap_sentences: int = 0) -> List[TextChunk]:

        """

        基于句子切片

        

        Args:

            text: 输入文本

            max_sentences: 每个块的最大句子数

            overlap_sentences: 块之间的重叠句子数

        """

        sentences = sent_tokenize(text)

        chunks = []

        chunk_id = 0

        i = 0

        

        while i < len(sentences):

            start = max(0, i - overlap_sentences) if i > 0 else i

            end = min(i + max_sentences, len(sentences))

            

            chunk_text = ' '.join(sentences[start:end])

            chunks.append(TextChunk(

                id=chunk_id,

                text=chunk_text,

                metadata={

                    "strategy": "sentence",

                    "sentence_count": end - start,

                    "sentence_range": f"{start}-{end}"

                },

                start_char=text.find(sentences[start]),

                end_char=text.find(sentences[end-1]) + len(sentences[end-1])

            ))

            

            chunk_id += 1

            i += max_sentences

        

        return chunks

    

    def chunk_by_paragraphs(self, text: str, max_paragraphs: int = 3) -> List[TextChunk]:

        """

        基于段落切片

        """

        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

        chunks = []

        chunk_id = 0

        

        for i in range(0, len(paragraphs), max_paragraphs):

            chunk_paragraphs = paragraphs[i:i+max_paragraphs]

            chunk_text = '\n\n'.join(chunk_paragraphs)

            

            chunks.append(TextChunk(

                id=chunk_id,

                text=chunk_text,

                metadata={

                    "strategy": "paragraph",

                    "paragraph_count": len(chunk_paragraphs),

                    "paragraph_range": f"{i}-{i+len(chunk_paragraphs)}"

                },

                start_char=0,  # 简化处理

                end_char=0

            ))

            chunk_id += 1

        

        return chunks

    

    def chunk_recursive(self, text: str, chunk_size: int = 500,

                       chunk_overlap: int = 50) -> List[TextChunk]:

        """

        递归字符文本切片（推荐方法）

        优先在段落、句子、单词边界处切分

        """

        from langchain.text_splitter import RecursiveCharacterTextSplitter

        

        text_splitter = RecursiveCharacterTextSplitter(

            chunk_size=chunk_size,

            chunk_overlap=chunk_overlap,

            length_function=len,

            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""]

        )

        

        split_texts = text_splitter.split_text(text)

        

        chunks = []

        for i, chunk_text in enumerate(split_texts):

            chunks.append(TextChunk(

                id=i,

                text=chunk_text,

                metadata={

                    "strategy": "recursive",

                    "chunk_size": chunk_size,

                    "chunk_overlap": chunk_overlap

                },

                start_char=0,

                end_char=0

            ))

        

        return chunks

    

    def chunk_with_overlap(self, text: str, chunk_size: int = 500,

                          overlap_size: int = 50) -> List[TextChunk]:

        """

        带重叠的切片，保持上下文连贯性

        """

        chunks = []

        start = 0

        text_length = len(text)

        chunk_id = 0

        

        while start < text_length:

            end = min(start + chunk_size, text_length)

            

            # 确保在单词边界处切分

            if end < text_length:

                # 向后查找空格

                space_pos = text.rfind(' ', start, end)

                if space_pos != -1:

                    end = space_pos

            

            chunk_text = text[start:end].strip()

            if chunk_text:

                chunks.append(TextChunk(

                    id=chunk_id,

                    text=chunk_text,

                    metadata={

                        "strategy": "overlapping",

                        "overlap_size": overlap_size,

                        "position": f"{start}-{end}"

                    },

                    start_char=start,

                    end_char=end

                ))

                chunk_id += 1

            

            # 移动起始位置，考虑重叠

            start = end - overlap_size if overlap_size > 0 else end

        

        return chunks

    

    def chunk_semantic(self, text: str, max_tokens: int = 512) -> List[TextChunk]:

        """

        基础语义切片（基于句子聚类）

        将语义相近的句子组合在一起

        """

        sentences = sent_tokenize(text)

        chunks = []

        current_chunk = []

        current_length = 0

        

        # 简单的启发式：使用tokenizer估算长度

        try:

            enc = tiktoken.get_encoding("cl100k_base")

            

            for sentence in sentences:

                sentence_tokens = len(enc.encode(sentence))

                

                if current_length + sentence_tokens > max_tokens and current_chunk:

                    # 保存当前块

                    chunk_text = ' '.join(current_chunk)

                    chunks.append(TextChunk(

                        id=len(chunks),

                        text=chunk_text,

                        metadata={

                            "strategy": "semantic",

                            "max_tokens": max_tokens,

                            "actual_tokens": current_length

                        },

                        start_char=0,

                        end_char=0

                    ))

                    current_chunk = []

                    current_length = 0

                

                current_chunk.append(sentence)

                current_length += sentence_tokens

            

            # 最后一个块

            if current_chunk:

                chunk_text = ' '.join(current_chunk)

                chunks.append(TextChunk(

                    id=len(chunks),

                    text=chunk_text,

                    metadata={

                        "strategy": "semantic",

                        "max_tokens": max_tokens,

                        "actual_tokens": current_length

                    },

                    start_char=0,

                    end_char=0

                ))

        except:

            # 降级到固定大小切片

            chunks = self.chunk_fixed_size(text, 500)

        

        return chunks

    

    def chunk(self, text: str, **kwargs) -> List[TextChunk]:

        """统一切片接口"""

        if self.strategy == ChunkingStrategy.FIXED_SIZE:

            return self.chunk_fixed_size(text, **kwargs)

        elif self.strategy == ChunkingStrategy.SENTENCE:

            return self.chunk_by_sentences(text, **kwargs)

        elif self.strategy == ChunkingStrategy.PARAGRAPH:

            return self.chunk_by_paragraphs(text, **kwargs)

        elif self.strategy == ChunkingStrategy.RECURSIVE:

            return self.chunk_recursive(text, **kwargs)

        elif self.strategy == ChunkingStrategy.SEMANTIC:

            return self.chunk_semantic(text, **kwargs)

        elif self.strategy == ChunkingStrategy.OVERLAPPING:

            return self.chunk_with_overlap(text, **kwargs)

        else:

            raise ValueError(f"不支持的切片策略: {self.strategy}")


class ChunkAnalyzer:

    """切片质量分析器"""

    @staticmethod

    def analyze_chunks(chunks: List[TextChunk]) -> Dict:

        """分析切片统计信息"""

        if not chunks:

            return {}

      

        lengths = [len(chunk.text) for chunk in chunks]

        word_counts = [len(chunk.text.split()) for chunk in chunks]

        

        analysis = {

            "total_chunks": len(chunks),

            "total_chars": sum(lengths),

            "total_words": sum(word_counts),

            "avg_chunk_length": sum(lengths) / len(chunks),

            "std_chunk_length": np.std(lengths) if len(lengths) > 1 else 0,

            "min_chunk_length": min(lengths),

            "max_chunk_length": max(lengths),

            "avg_words_per_chunk": sum(word_counts) / len(chunks),

            "chunk_length_distribution": {

                "<100": sum(1 for l in lengths if l < 100),

                "100-300": sum(1 for l in lengths if 100 <= l < 300),

                "300-500": sum(1 for l in lengths if 300 <= l < 500),

                "500-1000": sum(1 for l in lengths if 500 <= l < 1000),

                ">1000": sum(1 for l in lengths if l >= 1000)

            }

        }

        

        return analysis

    

    @staticmethod

    def print_report(analysis: Dict):

        """打印分析报告"""

        print("\n" + "="*60)

        print("文本切片分析报告")

        print("="*60)

        print(f"总切片数: {analysis['total_chunks']}")

        print(f"总字符数: {analysis['total_chars']:,}")

        print(f"总词数: {analysis['total_words']:,}")

        print(f"平均切片长度: {analysis['avg_chunk_length']:.1f} 字符")

        print(f"切片长度标准差: {analysis['std_chunk_length']:.1f}")

        print(f"最小切片长度: {analysis['min_chunk_length']} 字符")

        print(f"最大切片长度: {analysis['max_chunk_length']} 字符")

        print(f"平均词数/切片: {analysis['avg_words_per_chunk']:.1f}")

        print("\n切片长度分布:")

        for range_name, count in analysis['chunk_length_distribution'].items():

            percentage = count / analysis['total_chunks'] * 100

            print(f"  {range_name}: {count} ({percentage:.1f}%)")

        print("="*60 + "\n")

    

    @staticmethod

    def save_chunks_to_file(chunks: List[TextChunk], output_path: str):

        """保存切片到文件"""

        with open(output_path, 'w', encoding='utf-8') as f:

            for chunk in chunks:

                f.write(f"{'='*60}\n")

                f.write(f"Chunk ID: {chunk.id}\n")

                f.write(f"Metadata: {json.dumps(chunk.metadata, ensure_ascii=False)}\n")

                f.write(f"{'-'*60}\n")

                f.write(chunk.text)

                f.write(f"\n{'='*60}\n\n")